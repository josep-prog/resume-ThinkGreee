#!/usr/bin/env python
# coding: utf-8

# In[1]:


get_ipython().run_line_magic('pip', 'install PyPDF2 docx2txt docx')


# In[2]:


get_ipython().system('pip install python-docx')


# In[3]:


import os
import re
import json
import pandas as pd
import numpy as np
from pathlib import Path
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# NLP and ML libraries
import spacy
from transformers import (
    AutoTokenizer, AutoModelForTokenClassification,
    TrainingArguments, Trainer, pipeline
)
import torch
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.model_selection import train_test_split
from transformers import DataCollatorWithPadding

# Document processing
import PyPDF2
import docx
from io import BytesIO

# Data visualization
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots


# In[4]:


# Install additional required packages for improved image processing
get_ipython().run_line_magic('pip', 'install pytesseract Pillow easyocr')
print(" Additional image processing packages installed")
get_ipython().system('pip install PyMuPDF')


# In[5]:


import os
import re
import json
import pandas as pd
import numpy as np
from pathlib import Path
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# NLP and ML libraries
import spacy
from transformers import (
    AutoTokenizer, AutoModelForTokenClassification,
    TrainingArguments, Trainer, pipeline
)
import torch
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.model_selection import train_test_split

# Document processing
import PyPDF2
import docx
from io import BytesIO
from PIL import Image # Import PIL for image processing
import fitz # Import PyMuPDF

# Data visualization
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# Image processing and OCR
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️ pytesseract not available. Image text extraction will be limited.")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False
    print("⚠️ easyocr not available. Using alternative OCR methods.")

class ResumeParser:
    def __init__(self, model_name="AventIQ-AI/Resume-Parsing-NER-AI-Model"):
        """
        Initialize the Resume Parser with pretrained model
        """
        self.model_name = model_name
        self.tokenizer = None
        self.model = None
        self.nlp_pipeline = None
        self.resume_data = []
        self.processed_resumes = []

        # Initialize OCR reader if available
        self.ocr_reader = None
        if EASYOCR_AVAILABLE:
            try:
                self.ocr_reader = easyocr.Reader(['en'], gpu=False)
                print("✅ EasyOCR initialized successfully")
            except Exception as e:
                print(f"⚠️ Error initializing EasyOCR: {e}")

        # Initialize the model
        self.load_model()

        # Skills database (expandable)
        self.skills_database = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'go', 'rust', 'swift', 'kotlin', 'php', 'ruby', 'scala', 'r'],
            'web_development': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring'],
            'data_science': ['pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'matplotlib', 'seaborn', 'plotly'],
            'databases': ['mysql', 'postgresql', 'mongodb', 'sqlite', 'redis', 'elasticsearch', 'cassandra', 'oracle'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'gitlab'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem-solving', 'analytical', 'creative', 'management']
        }

    def load_model(self):
        """Load the pretrained model and tokenizer"""
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            self.model = AutoModelForTokenClassification.from_pretrained(self.model_name)
            self.nlp_pipeline = pipeline(
                "ner",
                model=self.model,
                tokenizer=self.tokenizer,
                aggregation_strategy="simple"
            )
            print(f" Model {self.model_name} loaded successfully!")
        except Exception as e:
            print(f" Error loading model: {e}")
            print(" Loading backup spaCy model...")
            try:
                import spacy
                self.nlp = spacy.load("en_core_web_sm")
                print(" SpaCy model loaded as backup!")
            except:
                print(" Please install spaCy model: python -m spacy download en_core_web_sm")

    def preprocess_resume(self, file_path):
        """
        Preprocess resume files (PDF, DOCX, TXT)
        """
        try:
            file_extension = Path(file_path).suffix.lower()

            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
                images = self._extract_pdf_images(file_path)
                # Extract text from images using OCR
                image_text = self._extract_text_from_images(images)
                combined_text = self._clean_text(text) + "\n" + image_text
                return combined_text, images
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
                images = self._extract_docx_images(file_path)
                image_text = self._extract_text_from_images(images)
                combined_text = self._clean_text(text) + "\n" + image_text
                return combined_text, images
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
                return self._clean_text(text), []
            else:
                print(f"⚠️ Unsupported file format: {file_extension}")
                return None, []

        except Exception as e:
            print(f"❌ Error preprocessing {file_path}: {e}")
            return None, []

    def _extract_pdf_text(self, file_path):
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            print(f" Error extracting PDF text: {e}")
            return None

    def _extract_pdf_images(self, file_path):
        """Extract images from PDF files with improved handling"""
        images = []
        try:
            doc = fitz.open(file_path)
            for page_index in range(len(doc)):
                page = doc[page_index]
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # Convert bytes to PIL Image for processing
                    try:
                        pil_image = Image.open(BytesIO(image_bytes))

                        # Store comprehensive image data
                        images.append({
                            'bytes': image_bytes,
                            'extension': image_ext,
                            'page': page_index,
                            'index': img_index,
                            'size': pil_image.size,
                            'mode': pil_image.mode,
                            'pil_image': pil_image
                        })
                    except Exception as e:
                        print(f"⚠️ Error processing image {img_index} on page {page_index}: {e}")

            doc.close()
            print(f"✅ Extracted {len(images)} images from PDF")
            return images
        except Exception as e:
            print(f"❌ Error extracting PDF images: {e}")
            return []

    def _extract_docx_images(self, file_path):
        """Extract images from DOCX files"""
        images = []
        try:
            doc = docx.Document(file_path)

            # Extract images from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_bytes = rel.target_part.blob
                        image_ext = rel.target_ref.split('.')[-1]

                        # Convert to PIL Image
                        pil_image = Image.open(BytesIO(image_bytes))

                        images.append({
                            'bytes': image_bytes,
                            'extension': image_ext,
                            'page': 0,  # DOCX doesn't have pages
                            'size': pil_image.size,
                            'mode': pil_image.mode,
                            'pil_image': pil_image
                        })
                    except Exception as e:
                        print(f"⚠️ Error processing DOCX image: {e}")

            print(f"✅ Extracted {len(images)} images from DOCX")
            return images
        except Exception as e:
            print(f"❌ Error extracting DOCX images: {e}")
            return []

    def _extract_text_from_images(self, images):
        """Extract text from images using OCR"""
        if not images:
            return ""

        extracted_text = []

        for idx, img_data in enumerate(images):
            try:
                pil_image = img_data.get('pil_image')
                if pil_image is None:
                    continue

                # Try EasyOCR first (better accuracy)
                if self.ocr_reader:
                    try:
                        result = self.ocr_reader.readtext(np.array(pil_image), detail=0)
                        text = " ".join(result)
                        if text.strip():
                            extracted_text.append(text)
                            print(f"✅ Extracted text from image {idx+1} using EasyOCR")
                    except Exception as e:
                        print(f"⚠️ EasyOCR failed for image {idx+1}: {e}")

                # Fallback to Tesseract
                elif TESSERACT_AVAILABLE:
                    try:
                        text = pytesseract.image_to_string(pil_image)
                        if text.strip():
                            extracted_text.append(text)
                            print(f"✅ Extracted text from image {idx+1} using Tesseract")
                    except Exception as e:
                        print(f"⚠️ Tesseract failed for image {idx+1}: {e}")

            except Exception as e:
                print(f"❌ Error extracting text from image {idx+1}: {e}")

        return "\n".join(extracted_text)

    def _extract_docx_text(self, file_path):
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f" Error extracting DOCX text: {e}")
            return None

    def _extract_txt_text(self, file_path):
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text
        except Exception as e:
            print(f" Error extracting TXT text: {e}")
            return None


    def _clean_text(self, text):
        """Clean and normalize text"""
        if not text:
            return ""

        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()

        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s@.-]', '', text)

        return text

    def save_extracted_images(self, candidate_data, output_folder="extracted_images"):
        """Save extracted images to disk for review"""
        try:
            output_path = Path(output_folder)
            output_path.mkdir(exist_ok=True)

            candidate_name = candidate_data.get('name', 'unknown').replace(' ', '_')
            candidate_folder = output_path / candidate_name
            candidate_folder.mkdir(exist_ok=True)

            images = candidate_data.get('images', [])
            saved_paths = []

            for idx, img_data in enumerate(images):
                try:
                    pil_image = img_data.get('pil_image')
                    ext = img_data.get('extension', 'png')

                    if pil_image:
                        image_path = candidate_folder / f"image_{idx+1}.{ext}"
                        pil_image.save(image_path)
                        saved_paths.append(str(image_path))

                except Exception as e:
                    print(f"⚠️ Error saving image {idx+1}: {e}")

            if saved_paths:
                print(f"✅ Saved {len(saved_paths)} images for {candidate_name}")

            return saved_paths

        except Exception as e:
            print(f"❌ Error saving images: {e}")
            return []

    def extract_candidate_data(self, text, images, filename):
        """
        Extract structured data from resume text using NER model
        """
        if not text:
            return None

        try:
            # Use NER model if available, otherwise use regex patterns
            if self.nlp_pipeline:
                entities = self.nlp_pipeline(text)
                candidate_data = self._process_ner_entities(entities, text)
            else:
                candidate_data = self._extract_with_regex(text)

           # Add metadata
            candidate_data['filename'] = filename
            candidate_data['processed_date'] = datetime.now().isoformat()
            candidate_data['raw_text'] = text

            # Enhanced image metadata
            candidate_data['images'] = []
            candidate_data['image_count'] = len(images)
            candidate_data['has_profile_picture'] = len(images) > 0

            for img in images:
                # Store only serializable image metadata
                img_metadata = {
                    'extension': img.get('extension'),
                    'page': img.get('page'),
                    'size': img.get('size'),
                    'mode': img.get('mode')
                }
                candidate_data['images'].append(img_metadata)

            return candidate_data

        except Exception as e:
            print(f" Error extracting candidate data: {e}")
            return self._extract_with_regex(text, filename)

    def _process_ner_entities(self, entities, text):
        """Process NER entities into structured data"""
        candidate_data = {
            'name': '',
            'email': '',
            'phone': '',
            'skills': [],
            'experience': [],
            'education': [],
            'organizations': [],
            'locations': []
        }

        # Process entities from NER model
        for entity in entities:
            label = entity['entity_group'].lower()
            value = entity['word'].strip()

            if 'person' in label or 'name' in label:
                if not candidate_data['name']:
                    candidate_data['name'] = value
            elif 'org' in label:
                candidate_data['organizations'].append(value)
            elif 'loc' in label:
                candidate_data['locations'].append(value)

        # Extract additional information using regex
        additional_data = self._extract_with_regex(text)

        # Merge data
        for key, value in additional_data.items():
            if key in candidate_data and value:
                if isinstance(candidate_data[key], list):
                    candidate_data[key].extend(value if isinstance(value, list) else [value])
                elif not candidate_data[key]:
                    candidate_data[key] = value

        return candidate_data

    def _extract_with_regex(self, text, filename=""):
        """Extract information using regex patterns"""
        candidate_data = {
            'filename': filename,
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'skills': self._extract_skills(text),
            'experience': self._extract_experience(text),
            'education': self._extract_education(text),
            'organizations': [],
            'locations': []
        }

        return candidate_data

    def _extract_name(self, text):
        """Extract candidate name"""
        # Look for name patterns at the beginning of resume
        lines = text.split('\n')[:5]  # Check first 5 lines

        for line in lines:
            line = line.strip()
            if len(line.split()) >= 2 and len(line) < 50:
                # Check if it looks like a name (no numbers, not email, not phone)
                if not re.search(r'[\d@]', line) and not re.search(r'(resume|cv|curriculum)', line.lower()):
                    return line

        return ""

    def _extract_email(self, text):
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else ""

    def _extract_phone(self, text):
        """Extract phone number"""
        phone_patterns = [
            r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'(\+\d{1,3}[-.\s]?)?\d{10,}',
        ]

        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]

        return ""

    def _extract_skills(self, text):
        """Extract skills from resume text"""
        text_lower = text.lower()
        found_skills = []

        # Check against skills database
        for category, skills in self.skills_database.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    found_skills.append(skill)

        # Look for skills sections
        skills_section_pattern = r'(?i)(skills?|technical skills?|core competencies|technologies)(.*?)(?=\n\s*\n|\n[A-Z]|\Z)'
        skills_match = re.search(skills_section_pattern, text, re.DOTALL)

        if skills_match:
            skills_text = skills_match.group(2)
            # Extract comma-separated or bullet-pointed skills
            additional_skills = re.findall(r'[•\-\*]?\s*([A-Za-z][A-Za-z0-9+#\.\s]{2,20})(?=[,\n•\-\*]|$)', skills_text)
            found_skills.extend([skill.strip() for skill in additional_skills if len(skill.strip()) > 2])

        return list(set(found_skills))  # Remove duplicates

    def _extract_experience(self, text):
        """Extract work experience"""
        experience_patterns = [
            r'(?i)(experience|work history|employment)(.*?)(?=education|skills|\Z)',
            r'(\d{4}[-\s]*\d{4}|\d{4}[-\s]*present).*?([A-Za-z\s,]+?)(?=\n|\Z)'
        ]

        experiences = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    experiences.extend([m.strip() for m in match if m.strip()])
                else:
                    experiences.append(match.strip())

        return experiences[:5]  # Return top 5 experiences

    def _extract_education(self, text):
        """Extract education information"""
        education_patterns = [
            r'(?i)(bachelor|master|phd|doctorate|diploma|degree).*?(?=\n|,|\Z)',
            r'(?i)(university|college|institute).*?(?=\n|,|\Z)',
            r'\b(19|20)\d{2}\b.*?(bachelor|master|phd|b\.?s\.?|m\.?s\.?|b\.?a\.?|m\.?a\.?)'
        ]

        education = []
        for pattern in education_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education.extend([match if isinstance(match, str) else ' '.join(match) for match in matches])

        return education[:3]  # Return top 3 education entries

    def process_resume_folder(self, folder_path):
        """
        Process all resumes in the specified folder
        """
        folder_path = Path(folder_path)
        if not folder_path.exists():
            print(f" Folder not found: {folder_path}")
            return

        print(f" Processing resumes in: {folder_path}")

        # Supported file extensions
        supported_extensions = ['.pdf', '.docx', '.txt']
        resume_files = []

        for ext in supported_extensions:
            resume_files.extend(folder_path.glob(f'*{ext}'))

        if not resume_files:
            print(" No resume files found in the folder")
            return

        print(f" Found {len(resume_files)} resume files")

        # Process each resume
        for i, file_path in enumerate(resume_files, 1):
            print(f"Processing {i}/{len(resume_files)}: {file_path.name}")

            # Preprocess resume (extract text and images)
            text, images = self.preprocess_resume(file_path)
            if not text:
                continue

            # Extract candidate data
            candidate_data = self.extract_candidate_data(text, images, file_path.name)
            if candidate_data:
                self.resume_data.append(candidate_data)

        print(f" Successfully processed {len(self.resume_data)} resumes")

    def rank_candidates(self, job_description, top_k=10):
        """
        Rank candidates based on job description similarity
        """
        if not self.resume_data:
            print("No resume data available. Please process resumes first.")
            return []

        print(f" Ranking candidates against job description...")

        # Prepare texts for similarity calculation
        job_desc_text = job_description.lower()
        resume_texts = []
        candidate_scores = []

        for candidate in self.resume_data:
            # Combine all text fields for similarity calculation
            resume_text = f"{' '.join(candidate.get('skills', []))} {' '.join(candidate.get('experience', []))} {' '.join(candidate.get('education', []))}"
            resume_texts.append(resume_text.lower())

        # Calculate TF-IDF similarity
        if resume_texts:
            vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)

            # Fit vectorizer on all texts
            all_texts = [job_desc_text] + resume_texts
            tfidf_matrix = vectorizer.fit_transform(all_texts)

            # Calculate similarity scores
            job_vector = tfidf_matrix[0:1]
            resume_vectors = tfidf_matrix[1:]

            similarities = cosine_similarity(job_vector, resume_vectors)[0]

            # Add scores to candidates
            for i, candidate in enumerate(self.resume_data):
                candidate['similarity_score'] = float(similarities[i])
                candidate['skill_match_score'] = self._calculate_skill_match(candidate, job_desc_text)
                candidate['total_score'] = (candidate['similarity_score'] * 0.7 +
                                          candidate['skill_match_score'] * 0.3)

        # Sort by total score
        ranked_candidates = sorted(self.resume_data, key=lambda x: x.get('total_score', 0), reverse=True)

        return ranked_candidates[:top_k]

    def _calculate_skill_match(self, candidate, job_description):
        """Calculate skill match score"""
        candidate_skills = [skill.lower() for skill in candidate.get('skills', [])]
        job_skills = []

        # Extract skills mentioned in job description
        for category, skills in self.skills_database.items():
            for skill in skills:
                if skill.lower() in job_description:
                    job_skills.append(skill.lower())

        if not job_skills:
            return 0.0

        # Calculate overlap
        matched_skills = set(candidate_skills) & set(job_skills)
        skill_match_score = len(matched_skills) / len(job_skills)

        return skill_match_score

    def generate_ats_report(self, ranked_candidates):
        """
        Generate ATS-style report with candidate rankings
        """
        print(" Generating ATS Report...")

        report_data = []
        for i, candidate in enumerate(ranked_candidates, 1):
            report_data.append({
                'Rank': i,
                'Name': candidate.get('name', 'N/A'),
                'Email': candidate.get('email', 'N/A'),
                'Phone': candidate.get('phone', 'N/A'),
                'Total Score': round(candidate.get('total_score', 0), 3),
                'Similarity Score': round(candidate.get('similarity_score', 0), 3),
                'Skill Match Score': round(candidate.get('skill_match_score', 0), 3),
                'Skills Count': len(candidate.get('skills', [])),
                'Top Skills': ', '.join(candidate.get('skills', [])[:5]),
                'Filename': candidate.get('filename', 'N/A'),
                'Image Count': len(candidate.get('images', [])) # Add image count
            })

        df = pd.DataFrame(report_data)

        # Save to CSV
        output_file = f"ats_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        df.to_csv(output_file, index=False)
        print(f" ATS Report saved to: {output_file}")

        return df

    def create_analytics_dashboard(self, candidates_df):
        """
        Create analytics and visualization dashboard
        """
        print(" Creating Analytics Dashboard...")

        # Create subplots
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=(
                'Candidate Score Distribution',
                'Skills Frequency',
                'Top 10 Candidates',
                'Score Components Comparison'
            ),
            specs=[[{"secondary_y": False}, {"secondary_y": False}],
                   [{"secondary_y": False}, {"secondary_y": False}]]
        )

        # 1. Score Distribution
        fig.add_trace(
            go.Histogram(x=candidates_df['Total Score'], nbinsx=20, name='Total Score'),
            row=1, col=1
        )

        # 2. Skills Analysis (if we have skills data)
        if 'Skills Count' in candidates_df.columns:
            fig.add_trace(
                go.Bar(x=candidates_df['Name'][:10], y=candidates_df['Skills Count'][:10], name='Skills Count'),
                row=1, col=2
            )

        # 3. Top 10 Candidates
        top_10 = candidates_df.head(10)
        fig.add_trace(
            go.Bar(x=top_10['Name'], y=top_10['Total Score'], name='Top 10 Scores'),
            row=2, col=1
        )

        # 4. Score Components
        fig.add_trace(
            go.Scatter(
                x=candidates_df['Similarity Score'][:20],
                y=candidates_df['Skill Match Score'][:20],
                mode='markers',
                text=candidates_df['Name'][:20],
                name='Score Components'
            ),
            row=2, col=2
        )

        # Update layout
        fig.update_layout(
            title_text="Resume Analytics Dashboard",
            showlegend=False,
            height=800
        )

        # Save dashboard
        dashboard_file = f"analytics_dashboard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        fig.write_html(dashboard_file)
        print(f" Analytics Dashboard saved to: {dashboard_file}")

        return fig

    def export_training_data(self, output_file="training_data.json"):
        """
        Export training data with labels for manual review/editing
        """
        if not self.resume_data:
            print("❌ No resume data to export")
            return False

        print(f"📤 Exporting training data to {output_file}...")

        training_texts, training_labels = self._prepare_training_data()

        if not hasattr(self, 'label_map'):
            print("❌ Label map not initialized. Run _prepare_training_data first.")
            return False

        export_data = []
        for i, (words, labels) in enumerate(zip(training_texts, training_labels)):
            # Convert label IDs back to label names
            label_names = []
            for label_id in labels:
                # Find label name from ID
                label_name = next((k for k, v in self.label_map.items() if v == label_id), 'O')
                label_names.append(label_name)

            candidate_idx = min(i, len(self.resume_data) - 1)
            export_data.append({
                'id': i,
                'filename': self.resume_data[candidate_idx].get('filename', 'unknown'),
                'words': words,
                'labels': label_names,
                'label_ids': labels
            })

        # Save to JSON
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)

        print(f"✅ Exported {len(export_data)} training examples")
        print(f"💡 You can manually review and edit labels in: {output_file}")
        return True

    def load_training_data(self, input_file="training_data.json"):
        """
        Load manually edited training data
        """
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                data = json.load(f)

            training_texts = []
            training_labels = []

            for item in data:
                training_texts.append(item['words'])
                training_labels.append(item['label_ids'])

            print(f"✅ Loaded {len(training_texts)} training examples from {input_file}")
            return training_texts, training_labels

        except Exception as e:
            print(f"❌ Error loading training data: {e}")
            return None, None

    def retrain_model(self, training_data_path=None):
        """
        Retrain the NER model with new resume data using proper NER labels

        Args:
            training_data_path: Optional path to JSON file with manually labeled data
        """
        print(" Starting improved model retraining...")

        if not self.model or not self.tokenizer:
            print(" Base model not loaded. Cannot retrain.")
            return False

        try:
            # Load training data - either from file or generate from processed resumes
            if training_data_path:
                print(f" Loading training data from: {training_data_path}")
                training_texts, training_labels = self.load_training_data(training_data_path)
                if training_texts is None:
                    print(" Failed to load training data. Falling back to auto-generated labels...")
                    training_texts, training_labels = self._prepare_training_data()
            else:
                print(" Generating training data from processed resumes...")
                training_texts, training_labels = self._prepare_training_data()

            if not training_texts:
                print(" No training data available")
                return False

            print(f" Training on {len(training_texts)} resume examples")

            # Split data (80% train, 20% validation)
            train_texts, val_texts, train_labels, val_labels = train_test_split(
                training_texts, training_labels, test_size=0.2, random_state=42
            )

            print(f" Training set: {len(train_texts)} examples")
            print(f" Validation set: {len(val_texts)} examples")

            # Tokenize and align labels for NER
            def tokenize_and_align_labels(texts, labels):
                tokenized_inputs = self.tokenizer(
                    texts,
                    truncation=True,
                    padding='max_length',  # Changed to max_length for consistency
                    max_length=512,
                    is_split_into_words=True  # Important for NER
                )

                aligned_labels = []
                for i, label in enumerate(labels):
                    word_ids = tokenized_inputs.word_ids(batch_index=i)
                    previous_word_idx = None
                    label_ids = []

                    for word_idx in word_ids:
                        # Special tokens get -100 (ignored by PyTorch loss)
                        if word_idx is None:
                            label_ids.append(-100)
                        # Only label the first token of each word
                        elif word_idx != previous_word_idx:
                            label_ids.append(label[word_idx])
                        else:
                            label_ids.append(-100)
                        previous_word_idx = word_idx

                    aligned_labels.append(label_ids)

                tokenized_inputs["labels"] = aligned_labels
                return tokenized_inputs

            # Process datasets
            train_encodings = tokenize_and_align_labels(train_texts, train_labels)
            val_encodings = tokenize_and_align_labels(val_texts, val_labels)

            # Create dataset class
            class ResumeDataset(torch.utils.data.Dataset):
                def __init__(self, encodings):
                    self.encodings = encodings

                def __getitem__(self, idx):
                    item = {key: torch.tensor(val[idx]) for key, val in self.encodings.items()}
                    return item

                def __len__(self):
                    return len(self.encodings["input_ids"])

            train_dataset = ResumeDataset(train_encodings)
            val_dataset = ResumeDataset(val_encodings)

            # Training arguments
            training_args = TrainingArguments(
                output_dir='./improved_retrained_model',
                num_train_epochs=5,
                per_device_train_batch_size=8,
                per_device_eval_batch_size=8,
                learning_rate=3e-5,  # Optimized learning rate
                warmup_ratio=0.1,  # Gradual warmup
                weight_decay=0.01,
                logging_dir='./logs',
                logging_steps=10,
                eval_strategy="epoch",  # Evaluate after each epoch
                save_strategy="epoch",  # Save after each epoch
                load_best_model_at_end=True,  # Load best model
                metric_for_best_model="eval_loss",  # Use loss as metric
                save_total_limit=2,
                report_to="none",  # Disable wandb/tensorboard
            )

            # Initialize trainer with improved setup
            trainer = Trainer(
                model=self.model,
                args=training_args,
                train_dataset=train_dataset,
                eval_dataset=val_dataset,
            )

            # Train model
            print("\n Starting training...")
            print("=" * 60)
            train_result = trainer.train()

            # Save retrained model
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            model_save_path = f"./best_resume_model_{timestamp}"
            trainer.save_model(model_save_path)
            self.tokenizer.save_pretrained(model_save_path)

            print("\n" + "=" * 60)
            print(" Training completed successfully!")
            print(f" Final training loss: {train_result.training_loss:.4f}")
            print(f" Model saved to: {model_save_path}")
            print("=" * 60)

            # Optional: Reload the trained model
            print("\n Reloading trained model into parser...")
            self.model = AutoModelForTokenClassification.from_pretrained(model_save_path)
            self.tokenizer = AutoTokenizer.from_pretrained(model_save_path)
            self.nlp_pipeline = pipeline(
                "ner",
                model=self.model,
                tokenizer=self.tokenizer,
                aggregation_strategy="simple"
            )
            print("✅ Trained model loaded successfully!")

            return True

        except Exception as e:
            print(f"\n Error during retraining: {e}")
            import traceback
            traceback.print_exc()
            return False


    def _prepare_training_data(self):
        """Prepare training data from processed resumes with proper NER labels"""
        training_texts = []
        training_labels = []

        # Label mapping for NER
        self.label_map = {
            'O': 0,           # Outside
            'B-PER': 1,       # Beginning of Person
            'I-PER': 2,       # Inside Person
            'B-ORG': 3,       # Beginning of Organization
            'I-ORG': 4,       # Inside Organization
            'B-SKILL': 5,     # Beginning of Skill
            'I-SKILL': 6,     # Inside Skill
            'B-EMAIL': 7,     # Beginning of Email
            'I-EMAIL': 8,     # Inside Email
            'B-PHONE': 9,     # Beginning of Phone
            'I-PHONE': 10,    # Inside Phone
            'B-EDU': 11,      # Beginning of Education
            'I-EDU': 12,      # Inside Education
            'B-DATE': 13,     # Beginning of Date
            'I-DATE': 14,     # Inside Date
            'B-LOC': 15,      # Beginning of Location
            'I-LOC': 16       # Inside Location
        }

        print(f" Creating training labels for {len(self.resume_data)} resumes...")

        for idx, candidate in enumerate(self.resume_data):
            text = candidate.get('raw_text', '')
            if not text:
                continue

            # Split into words (tokens)
            words = text.split()
            if not words or len(words) < 3:
                continue

            # Create BIO labels for each word
            labels = self._create_bio_labels(words, candidate, text)

            if len(labels) == len(words):
                training_texts.append(words)
                training_labels.append(labels)

                if (idx + 1) % 10 == 0:
                    print(f"  Processed {idx + 1}/{len(self.resume_data)} resumes")

        print(f" Created {len(training_texts)} training examples with proper NER labels")
        return training_texts, training_labels

    def _create_bio_labels(self, words, candidate, full_text):
        """
        Create BIO (Begin-Inside-Outside) labels for each word
        """
        labels = [self.label_map['O']] * len(words)  # Default all to Outside
        text_lower = full_text.lower()

        # Helper function to mark entity spans
        def mark_entity(entity_text, label_prefix):
            if not entity_text:
                return

            entity_words = str(entity_text).lower().split()
            if not entity_words:
                return

            # Find the entity in the words list
            for i in range(len(words) - len(entity_words) + 1):
                # Check if we have a match
                match = True
                for j, entity_word in enumerate(entity_words):
                    if entity_word not in words[i + j].lower():
                        match = False
                        break

                if match:
                    # Mark the first word as B- (Beginning)
                    labels[i] = self.label_map[f'B-{label_prefix}']
                    # Mark subsequent words as I- (Inside)
                    for j in range(1, len(entity_words)):
                        if i + j < len(labels):
                            labels[i + j] = self.label_map[f'I-{label_prefix}']
                    break

        # 1. Mark name (PERSON)
        name = candidate.get('name', '')
        if name and len(name) > 2:
            mark_entity(name, 'PER')

        # 2. Mark email
        email = candidate.get('email', '')
        if email:
            mark_entity(email, 'EMAIL')

        # 3. Mark phone
        phone = candidate.get('phone', '')
        if phone:
            # Clean phone for matching
            phone_clean = ''.join(filter(str.isdigit, str(phone)))
            if len(phone_clean) >= 10:
                for i, word in enumerate(words):
                    word_digits = ''.join(filter(str.isdigit, word))
                    if len(word_digits) >= 10 and word_digits in phone_clean:
                        labels[i] = self.label_map['B-PHONE']

        # 4. Mark skills
        skills = candidate.get('skills', [])
        for skill in skills:
            if skill and len(skill) > 1:
                mark_entity(skill, 'SKILL')

        # 5. Mark organizations
        organizations = candidate.get('organizations', [])
        for org in organizations:
            if org and len(org) > 2:
                mark_entity(org, 'ORG')

        # 6. Mark education
        education = candidate.get('education', [])
        for edu in education:
            if edu and len(str(edu)) > 3:
                mark_entity(edu, 'EDU')

        # 7. Mark locations
        locations = candidate.get('locations', [])
        for loc in locations:
            if loc and len(loc) > 2:
                mark_entity(loc, 'LOC')

        # 8. Mark dates (simple pattern matching)
        import re
        date_patterns = [
            r'\b\d{4}\b',  # Year
            r'\b\d{4}-\d{4}\b',  # Year range
            r'\b\d{4}-present\b',  # Year to present
        ]
        for i, word in enumerate(words):
            for pattern in date_patterns:
                if re.search(pattern, word.lower()):
                    labels[i] = self.label_map['B-DATE']
                    break

        return labels

    def test_retrained_model(self, test_resume_path=None):
        """
        Test the retrained model on a sample resume
        """
        if not self.nlp_pipeline:
            print(" No model loaded for testing")
            return None

        print("\n Testing retrained model...")
        print("=" * 60)

        # Use a processed resume if no test path provided
        if test_resume_path:
            text, _ = self.preprocess_resume(test_resume_path)
        elif self.resume_data:
            # Use first resume from processed data
            text = self.resume_data[0].get('raw_text', '')
            print(f" Testing on: {self.resume_data[0].get('filename', 'Unknown')}")
        else:
            print(" No test data available")
            return None

        if not text:
            print(" No text to test")
            return None

        # Extract entities using the model
        entities = self.nlp_pipeline(text[:1000])  # Test on first 1000 chars

        print(f"\n Found {len(entities)} entities:\n")

        # Group entities by type
        entity_groups = {}
        for entity in entities:
            entity_type = entity['entity_group']
            if entity_type not in entity_groups:
                entity_groups[entity_type] = []
            entity_groups[entity_type].append({
                'text': entity['word'],
                'score': entity['score']
            })

        # Display results
        for entity_type, items in entity_groups.items():
            print(f"\n{entity_type}:")
            for item in items[:5]:  # Show first 5 of each type
                print(f"  • {item['text']} (confidence: {item['score']:.3f})")

        print("\n" + "=" * 60)
        return entities

    def save_processed_data(self):
        """Save all processed resume data"""
        if not self.resume_data:
            print(" No data to save")
            return

        # Create a copy without PIL images for JSON serialization
        serializable_data = []
        for candidate in self.resume_data:
            candidate_copy = candidate.copy()

            # Remove non-serializable PIL image objects if they exist
            if 'images' in candidate_copy:
                candidate_copy['images'] = [
                    {k: v for k, v in img.items() if k != 'pil_image'}
                    for img in candidate_copy.get('images', [])
                ]
            serializable_data.append(candidate_copy)

        # Save as JSON
        json_file = f"processed_resumes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(serializable_data, f, indent=2, ensure_ascii=False)

        print(f" Processed data saved to: {json_file}")


# In[7]:


from google.colab import drive
drive.mount('/content/drive')


# In[8]:


def main():
    """
    Main function to demonstrate the complete resume parsing pipeline
    """
    print(" Starting Advanced Resume Parser and ATS System")
    print("=" * 60)

    # Initialize parser
    parser = ResumeParser()

    # Process resumes from folder
    resume_folder = "/content/drive/MyDrive/Colab Notebooks/THINKGREEN AI/RESUMES"
    parser.process_resume_folder(resume_folder)

    if not parser.resume_data:
        print(" No resumes processed. Exiting.")
        return

    # Example job description
    job_description = """
    We are looking for a Python Developer with experience in machine learning,
    data science, and web development. Required skills include Python, pandas,
    scikit-learn, Django, REST APIs, and SQL. Experience with cloud platforms
    like AWS is preferred. Strong problem-solving and communication skills required.
    """

    print("\n Job Description:")
    print(job_description)
    print("=" * 60)

    # Rank candidates
    top_candidates = parser.rank_candidates(job_description, top_k=10)

    # Generate ATS report
    report_df = parser.generate_ats_report(top_candidates)

    # Display top candidates
    print("\n🏆 TOP CANDIDATES:")
    print("=" * 60)
    for i, candidate in enumerate(top_candidates[:5], 1):
        print(f"{i}. {candidate.get('name', 'N/A')} (Score: {candidate.get('total_score', 0):.3f})")
        print(f"    {candidate.get('email', 'N/A')}")
        print(f"    {candidate.get('phone', 'N/A')}")
        print(f"    Skills: {', '.join(candidate.get('skills', [])[:5])}")
        print(f"   🖼️  Images: {candidate.get('image_count', 0)}")
        print(f"   👤 Profile Picture: {'Yes' if candidate.get('has_profile_picture', False) else 'No'}")
        print(f"    File: {candidate.get('filename', 'N/A')}")
        print("-" * 40)

    # Create analytics dashboard
    dashboard = parser.create_analytics_dashboard(report_df)

    # Save processed data
    parser.save_processed_data()

    # Retrain model (optional)
    print("\n Retraining model with processed data...")
    parser.retrain_model()

    print("\n Complete Resume Parser and ATS System finished successfully!")
    print(" Check the generated files for detailed analysis.")

if __name__ == "__main__":
    main()


# In[ ]:





# # Visual Training Flow Diagram
# 
# ```
# ┌─────────────────────────────────────────────────────────────────┐
# │                    IMPROVED RETRAINING FLOW                     │
# └─────────────────────────────────────────────────────────────────┘
# 
# 📂 STEP 1: Process Resumes
#    ├── Extract Text (PDF/DOCX/TXT)
#    ├── Extract Images & OCR
#    ├── Extract Entities (Name, Email, Skills, etc.)
#    └── Store in resume_data[]
#         │
#         ▼
# ┌───────────────────────────────────────────────┐
# │  Raw Resume:                                  │
# │  "John Smith                                  │
# │   Email: john@email.com                       │
# │   Skills: Python, TensorFlow, AWS"            │
# │                                               │
# │  Extracted Data:                              │
# │  - name: "John Smith"                         │
# │  - email: "john@email.com"                    │
# │  - skills: ["Python", "TensorFlow", "AWS"]    │
# └───────────────────────────────────────────────┘
#         │
#         ▼
# 🏷️  STEP 2: Create BIO Labels (_create_bio_labels)
#    ├── Split text into words/tokens
#    ├── Match extracted name → B-PER, I-PER
#    ├── Match extracted email → B-EMAIL
#    ├── Match extracted skills → B-SKILL, I-SKILL
#    ├── Match organizations → B-ORG, I-ORG
#    └── Everything else → O (Outside)
#         │
#         ▼
# ┌──────────────────────────────────────────────────┐
# │  Words:  ["John", "Smith", "Python", "AWS"]      │
# │  Labels: [1,      2,       5,        5]          │
# │           ↓       ↓        ↓         ↓           │
# │        B-PER   I-PER   B-SKILL  B-SKILL          │
# └──────────────────────────────────────────────────┘
#         │
#         ▼
# STEP 3: Prepare Training Data
#    ├── Convert to tokenizer format
#    ├── Align labels with sub-word tokens
#    ├── Add special tokens ([CLS], [SEP])
#    └── Create PyTorch datasets
#         │
#         ▼
# ┌──────────────────────────────────────────────┐
# │  Tokenized Input:                            │
# │  [CLS] John Smith Python [SEP] [PAD] ...     │
# │  -100  1    2     5      -100  -100  ...     │
# │   ↓    ↓    ↓     ↓       ↓    ↓             │
# │  Skip B-PER I-PER B-SKILL Skip Skip           │
# └──────────────────────────────────────────────┘
#         │
#         ▼
# 🎓 STEP 4: Train Model (retrain_model)
#    ├── Split: 80% train, 20% validation
#    ├── Train for 5 epochs
#    ├── Evaluate after each epoch
#    ├── Save best model
#    └── Reload into pipeline
#         │
#         ▼
# ┌────────────────────────────────────────────┐
# │  Training Progress:                        │
# │  Epoch 1/5: Loss 2.456  ▓░░░░░░░░░░  10%   │
# │  Epoch 2/5: Loss 1.823  ▓▓▓░░░░░░░░  40%   │
# │  Epoch 3/5: Loss 1.234  ▓▓▓▓▓▓░░░░░  60%   │
# │  Epoch 4/5: Loss 0.892  ▓▓▓▓▓▓▓▓░░░  80%   │
# │  Epoch 5/5: Loss 0.645  ▓▓▓▓▓▓▓▓▓▓ 100%   │
# └────────────────────────────────────────────┘
#         │
#         ▼
#  STEP 5: Improved Model Ready!
#    ├── Better entity recognition
#    ├── Custom to your resume format
#    ├── Learns domain-specific terms
#    └── Higher confidence scores
# 
# ┌────────────────────────────────────────────┐
# │  Test on New Resume:                       │
# │  "Sarah Johnson, ML Engineer at Tesla"     │
# │                                            │
# │  Extracted:                                │
# │  ✅ PERSON: Sarah Johnson (0.98)          │
# │  ✅ ORG: Tesla (0.95)                     │
# │  ✅ SKILL: ML (0.87)                      │
# └────────────────────────────────────────────┘
# ```
# 
# ---
# 
# ## 🔄 Optional: Manual Refinement Loop
# 
# ```
# parser.export_training_data("labels.json")
#          │
#          ▼
#    Edit labels.json
#    (Fix any mistakes)
#          │
#          ▼
# parser.retrain_model("labels.json")
#          │
#          ▼
# parser.test_retrained_model()
#          │
#          ▼
#     Better Results! 🎯
# ```
# 

# In[ ]:




