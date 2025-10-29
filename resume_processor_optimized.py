#!/usr/bin/env python3
"""
Resume Parser and ATS System
Professional resume screening and candidate ranking system.
Processes resumes (PDF, DOCX, TXT) and grades them against job criteria.
"""

import os
import re
import json
import warnings
from pathlib import Path
from datetime import datetime
from io import BytesIO

import pandas as pd
import numpy as np
from PIL import Image

# NLP and ML
import spacy
from transformers import (
    AutoTokenizer, 
    AutoModelForTokenClassification,
    TrainingArguments, 
    Trainer, 
    pipeline
)
import torch
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.model_selection import train_test_split

# Document processing
import PyPDF2
import docx
import fitz  # PyMuPDF

# Visualization
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# OCR (optional)
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

warnings.filterwarnings('ignore')


class ResumeParser:
    """Parse and analyze resumes with NER-based entity extraction."""
    
    def __init__(self, model_name="AventIQ-AI/Resume-Parsing-NER-AI-Model"):
        """Initialize parser with pretrained model."""
        self.model_name = model_name
        self.tokenizer = None
        self.model = None
        self.nlp_pipeline = None
        self.resume_data = []
        self.ocr_reader = None
        
        # Initialize OCR
        if EASYOCR_AVAILABLE:
            try:
                self.ocr_reader = easyocr.Reader(['en'], gpu=False)
            except Exception:
                pass
        
        # Load model
        self.load_model()
        
        # Skills database
        self.skills_database = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'go', 'rust', 'swift', 'kotlin', 'php', 'ruby', 'scala', 'r', 'typescript'],
            'web_development': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'fastapi'],
            'data_science': ['pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'matplotlib', 'seaborn', 'plotly', 'sql'],
            'databases': ['mysql', 'postgresql', 'mongodb', 'sqlite', 'redis', 'elasticsearch', 'cassandra', 'oracle', 'dynamodb'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'gitlab', 'ci/cd'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem-solving', 'analytical', 'creative', 'management']
        }
        
        # NER label mapping
        self.label_map = {
            'O': 0, 'B-PER': 1, 'I-PER': 2, 'B-ORG': 3, 'I-ORG': 4,
            'B-SKILL': 5, 'I-SKILL': 6, 'B-EMAIL': 7, 'I-EMAIL': 8,
            'B-PHONE': 9, 'I-PHONE': 10, 'B-EDU': 11, 'I-EDU': 12,
            'B-DATE': 13, 'I-DATE': 14, 'B-LOC': 15, 'I-LOC': 16
        }
    
    def load_model(self):
        """Load pretrained NER model."""
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            self.model = AutoModelForTokenClassification.from_pretrained(self.model_name)
            self.nlp_pipeline = pipeline(
                "ner",
                model=self.model,
                tokenizer=self.tokenizer,
                aggregation_strategy="simple"
            )
            print(f"Model loaded: {self.model_name}")
        except Exception as e:
            print(f"Error loading model: {e}. Using backup spaCy model.")
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except Exception:
                print("Install spaCy model: python -m spacy download en_core_web_sm")
    
    def preprocess_resume(self, file_path):
        """Extract text and images from resume file."""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
                images = self._extract_pdf_images(file_path)
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
                images = self._extract_docx_images(file_path)
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
                images = []
            else:
                print(f"Unsupported format: {file_extension}")
                return None, []
            
            # Extract text from images
            image_text = self._extract_text_from_images(images) if images else ""
            combined_text = self._clean_text(text) + "\n" + image_text
            
            return combined_text, images
            
        except Exception as e:
            print(f"Error preprocessing {file_path}: {e}")
            return None, []
    
    def _extract_pdf_text(self, file_path):
        """Extract text from PDF."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return "\n".join(page.extract_text() for page in pdf_reader.pages)
        except Exception as e:
            print(f"PDF text extraction error: {e}")
            return ""
    
    def _extract_pdf_images(self, file_path):
        """Extract images from PDF."""
        images = []
        try:
            doc = fitz.open(file_path)
            for page_index in range(len(doc)):
                page = doc[page_index]
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    
                    try:
                        pil_image = Image.open(BytesIO(base_image["image"]))
                        images.append({
                            'bytes': base_image["image"],
                            'extension': base_image["ext"],
                            'page': page_index,
                            'index': img_index,
                            'size': pil_image.size,
                            'pil_image': pil_image
                        })
                    except Exception:
                        pass
            doc.close()
            return images
        except Exception as e:
            print(f"PDF image extraction error: {e}")
            return []
    
    def _extract_docx_text(self, file_path):
        """Extract text from DOCX."""
        try:
            doc = docx.Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            print(f"DOCX text extraction error: {e}")
            return ""
    
    def _extract_docx_images(self, file_path):
        """Extract images from DOCX."""
        images = []
        try:
            doc = docx.Document(file_path)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_bytes = rel.target_part.blob
                        pil_image = Image.open(BytesIO(image_bytes))
                        images.append({
                            'bytes': image_bytes,
                            'extension': rel.target_ref.split('.')[-1],
                            'size': pil_image.size,
                            'pil_image': pil_image
                        })
                    except Exception:
                        pass
            return images
        except Exception as e:
            print(f"DOCX image extraction error: {e}")
            return []
    
    def _extract_txt_text(self, file_path):
        """Extract text from TXT."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"TXT extraction error: {e}")
            return ""
    
    def _extract_text_from_images(self, images):
        """Extract text from images using OCR."""
        if not images:
            return ""
        
        extracted_text = []
        for img_data in images:
            pil_image = img_data.get('pil_image')
            if not pil_image:
                continue
            
            # Try EasyOCR
            if self.ocr_reader:
                try:
                    result = self.ocr_reader.readtext(np.array(pil_image), detail=0)
                    text = " ".join(result)
                    if text.strip():
                        extracted_text.append(text)
                    continue
                except Exception:
                    pass
            
            # Try Tesseract
            if TESSERACT_AVAILABLE:
                try:
                    text = pytesseract.image_to_string(pil_image)
                    if text.strip():
                        extracted_text.append(text)
                except Exception:
                    pass
        
        return "\n".join(extracted_text)
    
    def _clean_text(self, text):
        """Clean and normalize text."""
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text).strip()
        text = re.sub(r'[^\w\s@.-]', '', text)
        return text
    
    def extract_candidate_data(self, text, images, filename):
        """Extract structured data from resume."""
        if not text:
            return None
        
        try:
            # Use NER pipeline or regex
            if self.nlp_pipeline:
                entities = self.nlp_pipeline(text)
                candidate_data = self._process_ner_entities(entities, text)
            else:
                candidate_data = self._extract_with_regex(text)
            
            # Add metadata
            candidate_data['filename'] = filename
            candidate_data['processed_date'] = datetime.now().isoformat()
            candidate_data['raw_text'] = text
            candidate_data['image_count'] = len(images)
            candidate_data['has_profile_picture'] = len(images) > 0
            
            # Store serializable image metadata
            candidate_data['images'] = [
                {
                    'extension': img.get('extension'),
                    'page': img.get('page'),
                    'size': img.get('size')
                }
                for img in images
            ]
            
            return candidate_data
            
        except Exception as e:
            print(f"Extraction error: {e}")
            return self._extract_with_regex(text, filename)
    
    def _process_ner_entities(self, entities, text):
        """Process NER entities into structured data."""
        candidate_data = {
            'name': '', 'email': '', 'phone': '',
            'skills': [], 'experience': [], 'education': [],
            'organizations': [], 'locations': []
        }
        
        for entity in entities:
            label = entity['entity_group'].lower()
            value = entity['word'].strip()
            
            if 'person' in label or 'name' in label:
                if not candidate_data['name']:
                    candidate_data['name'] = value
            elif 'org' in label:
                candidate_data['organizations'].append(value)
            elif 'loc' in label:
                candidate_data['locations'].append(value)
        
        # Merge with regex extraction
        additional_data = self._extract_with_regex(text)
        for key, value in additional_data.items():
            if key in candidate_data and value:
                if isinstance(candidate_data[key], list):
                    candidate_data[key].extend(value if isinstance(value, list) else [value])
                elif not candidate_data[key]:
                    candidate_data[key] = value
        
        return candidate_data
    
    def _extract_with_regex(self, text, filename=""):
        """Extract information using regex patterns."""
        return {
            'filename': filename,
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'skills': self._extract_skills(text),
            'experience': self._extract_experience(text),
            'education': self._extract_education(text),
            'organizations': [],
            'locations': []
        }
    
    def _extract_name(self, text):
        """Extract candidate name from first few lines."""
        lines = text.split('\n')[:5]
        for line in lines:
            line = line.strip()
            if 2 <= len(line.split()) <= 5 and len(line) < 50:
                if not re.search(r'[\d@]', line) and not re.search(r'(resume|cv|curriculum)', line.lower()):
                    return line
        return ""
    
    def _extract_email(self, text):
        """Extract email address."""
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(pattern, text)
        return emails[0] if emails else ""
    
    def _extract_phone(self, text):
        """Extract phone number."""
        patterns = [
            r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'(\+\d{1,3}[-.\s]?)?\d{10,}',
        ]
        for pattern in patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        return ""
    
    def _extract_skills(self, text):
        """Extract skills from resume."""
        text_lower = text.lower()
        found_skills = set()
        
        # Match against skills database
        for skills in self.skills_database.values():
            for skill in skills:
                if skill.lower() in text_lower:
                    found_skills.add(skill)
        
        # Extract from skills section
        skills_pattern = r'(?i)(skills?|technical skills?|core competencies)(.*?)(?=\n\s*\n|\n[A-Z]|\Z)'
        match = re.search(skills_pattern, text, re.DOTALL)
        if match:
            skills_text = match.group(2)
            additional = re.findall(r'[•\-\*]?\s*([A-Za-z][A-Za-z0-9+#\.\s]{2,20})(?=[,\n•\-\*]|$)', skills_text)
            found_skills.update(skill.strip() for skill in additional if len(skill.strip()) > 2)
        
        return list(found_skills)
    
    def _extract_experience(self, text):
        """Extract work experience."""
        patterns = [
            r'(?i)(experience|work history|employment)(.*?)(?=education|skills|\Z)',
            r'(\d{4}[-\s]*\d{4}|\d{4}[-\s]*present).*?([A-Za-z\s,]+?)(?=\n|\Z)'
        ]
        experiences = []
        for pattern in patterns:
            for match in re.findall(pattern, text, re.DOTALL | re.IGNORECASE):
                if isinstance(match, tuple):
                    experiences.extend([m.strip() for m in match if m.strip()])
                else:
                    experiences.append(match.strip())
        return experiences[:5]
    
    def _extract_education(self, text):
        """Extract education information."""
        patterns = [
            r'(?i)(bachelor|master|phd|doctorate|diploma|degree).*?(?=\n|,|\Z)',
            r'(?i)(university|college|institute).*?(?=\n|,|\Z)',
            r'\b(19|20)\d{2}\b.*?(bachelor|master|phd|b\.?s\.?|m\.?s\.?)'
        ]
        education = []
        for pattern in patterns:
            for match in re.findall(pattern, text, re.IGNORECASE):
                education.append(match if isinstance(match, str) else ' '.join(match))
        return education[:3]
    
    def process_resume_folder(self, folder_path):
        """Process all resumes in folder."""
        folder_path = Path(folder_path)
        if not folder_path.exists():
            print(f"Folder not found: {folder_path}")
            return
        
        print(f"Processing resumes in: {folder_path}")
        
        # Find all supported files
        resume_files = []
        for ext in ['.pdf', '.docx', '.txt']:
            resume_files.extend(folder_path.glob(f'*{ext}'))
        
        if not resume_files:
            print("No resume files found")
            return
        
        print(f"Found {len(resume_files)} resumes")
        
        # Process each resume
        for i, file_path in enumerate(resume_files, 1):
            print(f"Processing {i}/{len(resume_files)}: {file_path.name}")
            
            text, images = self.preprocess_resume(file_path)
            if text:
                candidate_data = self.extract_candidate_data(text, images, file_path.name)
                if candidate_data:
                    self.resume_data.append(candidate_data)
        
        print(f"Successfully processed {len(self.resume_data)} resumes")
    
    def rank_candidates(self, job_description, top_k=10):
        """Rank candidates based on job description similarity."""
        if not self.resume_data:
            print("No resume data available")
            return []
        
        print("Ranking candidates...")
        
        job_desc_lower = job_description.lower()
        resume_texts = []
        
        for candidate in self.resume_data:
            resume_text = f"{' '.join(candidate.get('skills', []))} " \
                         f"{' '.join(candidate.get('experience', []))} " \
                         f"{' '.join(candidate.get('education', []))}"
            resume_texts.append(resume_text.lower())
        
        # Calculate TF-IDF similarity
        if resume_texts:
            vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
            all_texts = [job_desc_lower] + resume_texts
            tfidf_matrix = vectorizer.fit_transform(all_texts)
            
            job_vector = tfidf_matrix[0:1]
            resume_vectors = tfidf_matrix[1:]
            similarities = cosine_similarity(job_vector, resume_vectors)[0]
            
            # Calculate scores
            for i, candidate in enumerate(self.resume_data):
                candidate['similarity_score'] = float(similarities[i])
                candidate['skill_match_score'] = self._calculate_skill_match(candidate, job_desc_lower)
                candidate['total_score'] = (candidate['similarity_score'] * 0.7 + 
                                           candidate['skill_match_score'] * 0.3)
        
        # Sort by score
        ranked = sorted(self.resume_data, key=lambda x: x.get('total_score', 0), reverse=True)
        return ranked[:top_k]
    
    def _calculate_skill_match(self, candidate, job_description):
        """Calculate skill match score."""
        candidate_skills = {skill.lower() for skill in candidate.get('skills', [])}
        job_skills = set()
        
        for skills in self.skills_database.values():
            for skill in skills:
                if skill.lower() in job_description:
                    job_skills.add(skill.lower())
        
        if not job_skills:
            return 0.0
        
        matched = candidate_skills & job_skills
        return len(matched) / len(job_skills)
    
    def generate_ats_report(self, ranked_candidates):
        """Generate ATS report with rankings."""
        print("Generating ATS report...")
        
        report_data = []
        for i, candidate in enumerate(ranked_candidates, 1):
            report_data.append({
                'Rank': i,
                'Name': candidate.get('name', 'N/A'),
                'Email': candidate.get('email', 'N/A'),
                'Phone': candidate.get('phone', 'N/A'),
                'Total Score': round(candidate.get('total_score', 0), 3),
                'Similarity Score': round(candidate.get('similarity_score', 0), 3),
                'Skill Match Score': round(candidate.get('skill_match_score', 0), 3),
                'Skills Count': len(candidate.get('skills', [])),
                'Top Skills': ', '.join(candidate.get('skills', [])[:5]),
                'Filename': candidate.get('filename', 'N/A'),
                'Image Count': candidate.get('image_count', 0)
            })
        
        df = pd.DataFrame(report_data)
        
        # Save report
        output_file = f"ats_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        df.to_csv(output_file, index=False)
        print(f"Report saved: {output_file}")
        
        return df
    
    def create_analytics_dashboard(self, candidates_df):
        """Create visualization dashboard."""
        print("Creating dashboard...")
        
        fig = make_subplots(
            rows=2, cols=2,
            subplot_titles=(
                'Score Distribution',
                'Skills Analysis',
                'Top 10 Candidates',
                'Score Components'
            )
        )
        
        # Score distribution
        fig.add_trace(
            go.Histogram(x=candidates_df['Total Score'], nbinsx=20, name='Total Score'),
            row=1, col=1
        )
        
        # Skills count
        fig.add_trace(
            go.Bar(x=candidates_df['Name'][:10], y=candidates_df['Skills Count'][:10]),
            row=1, col=2
        )
        
        # Top candidates
        top_10 = candidates_df.head(10)
        fig.add_trace(
            go.Bar(x=top_10['Name'], y=top_10['Total Score']),
            row=2, col=1
        )
        
        # Score components scatter
        fig.add_trace(
            go.Scatter(
                x=candidates_df['Similarity Score'][:20],
                y=candidates_df['Skill Match Score'][:20],
                mode='markers',
                text=candidates_df['Name'][:20]
            ),
            row=2, col=2
        )
        
        fig.update_layout(
            title_text="Resume Analytics Dashboard",
            showlegend=False,
            height=800
        )
        
        dashboard_file = f"analytics_dashboard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        fig.write_html(dashboard_file)
        print(f"Dashboard saved: {dashboard_file}")
        
        return fig
    
    def retrain_model(self, training_data_path=None):
        """Retrain NER model with processed resumes."""
        print("Starting model retraining...")
        
        if not self.model or not self.tokenizer:
            print("Base model not loaded")
            return False
        
        try:
            # Prepare training data
            if training_data_path:
                training_texts, training_labels = self._load_training_data(training_data_path)
            else:
                training_texts, training_labels = self._prepare_training_data()
            
            if not training_texts:
                print("No training data available")
                return False
            
            print(f"Training on {len(training_texts)} examples")
            
            # Split data
            train_texts, val_texts, train_labels, val_labels = train_test_split(
                training_texts, training_labels, test_size=0.2, random_state=42
            )
            
            # Tokenize and align labels
            def tokenize_and_align(texts, labels):
                tokenized = self.tokenizer(
                    texts,
                    truncation=True,
                    padding='max_length',
                    max_length=512,
                    is_split_into_words=True
                )
                
                aligned_labels = []
                for i, label in enumerate(labels):
                    word_ids = tokenized.word_ids(batch_index=i)
                    prev_word_idx = None
                    label_ids = []
                    
                    for word_idx in word_ids:
                        if word_idx is None:
                            label_ids.append(-100)
                        elif word_idx != prev_word_idx:
                            label_ids.append(label[word_idx])
                        else:
                            label_ids.append(-100)
                        prev_word_idx = word_idx
                    
                    aligned_labels.append(label_ids)
                
                tokenized["labels"] = aligned_labels
                return tokenized
            
            train_encodings = tokenize_and_align(train_texts, train_labels)
            val_encodings = tokenize_and_align(val_texts, val_labels)
            
            # Create datasets
            class ResumeDataset(torch.utils.data.Dataset):
                def __init__(self, encodings):
                    self.encodings = encodings
                
                def __getitem__(self, idx):
                    return {key: torch.tensor(val[idx]) for key, val in self.encodings.items()}
                
                def __len__(self):
                    return len(self.encodings["input_ids"])
            
            train_dataset = ResumeDataset(train_encodings)
            val_dataset = ResumeDataset(val_encodings)
            
            # Training arguments
            training_args = TrainingArguments(
                output_dir='./retrained_model',
                num_train_epochs=5,
                per_device_train_batch_size=8,
                per_device_eval_batch_size=8,
                learning_rate=3e-5,
                warmup_ratio=0.1,
                weight_decay=0.01,
                logging_steps=10,
                eval_strategy="epoch",
                save_strategy="epoch",
                load_best_model_at_end=True,
                save_total_limit=2,
                report_to="none"
            )
            
            # Train
            trainer = Trainer(
                model=self.model,
                args=training_args,
                train_dataset=train_dataset,
                eval_dataset=val_dataset
            )
            
            print("Training...")
            result = trainer.train()
            
            # Save model
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            model_path = f"./resume_model_{timestamp}"
            trainer.save_model(model_path)
            self.tokenizer.save_pretrained(model_path)
            
            print(f"Training complete. Loss: {result.training_loss:.4f}")
            print(f"Model saved: {model_path}")
            
            # Reload model
            self.model = AutoModelForTokenClassification.from_pretrained(model_path)
            self.tokenizer = AutoTokenizer.from_pretrained(model_path)
            self.nlp_pipeline = pipeline(
                "ner",
                model=self.model,
                tokenizer=self.tokenizer,
                aggregation_strategy="simple"
            )
            
            return True
            
        except Exception as e:
            print(f"Retraining error: {e}")
            return False
    
    def _prepare_training_data(self):
        """Prepare training data with BIO labels."""
        training_texts = []
        training_labels = []
        
        for candidate in self.resume_data:
            text = candidate.get('raw_text', '')
            if not text:
                continue
            
            words = text.split()
            if len(words) < 3:
                continue
            
            labels = self._create_bio_labels(words, candidate, text)
            if len(labels) == len(words):
                training_texts.append(words)
                training_labels.append(labels)
        
        return training_texts, training_labels
    
    def _create_bio_labels(self, words, candidate, full_text):
        """Create BIO labels for words."""
        labels = [self.label_map['O']] * len(words)
        
        def mark_entity(entity_text, prefix):
            if not entity_text:
                return
            entity_words = str(entity_text).lower().split()
            
            for i in range(len(words) - len(entity_words) + 1):
                match = all(ew in words[i + j].lower() 
                           for j, ew in enumerate(entity_words))
                if match:
                    labels[i] = self.label_map[f'B-{prefix}']
                    for j in range(1, len(entity_words)):
                        if i + j < len(labels):
                            labels[i + j] = self.label_map[f'I-{prefix}']
                    break
        
        # Mark entities
        mark_entity(candidate.get('name'), 'PER')
        mark_entity(candidate.get('email'), 'EMAIL')
        
        for skill in candidate.get('skills', []):
            mark_entity(skill, 'SKILL')
        
        for org in candidate.get('organizations', []):
            mark_entity(org, 'ORG')
        
        return labels
    
    def save_processed_data(self):
        """Save processed resume data as JSON."""
        if not self.resume_data:
            print("No data to save")
            return
        
        # Remove non-serializable data
        serializable_data = []
        for candidate in self.resume_data:
            candidate_copy = candidate.copy()
            if 'images' in candidate_copy:
                candidate_copy['images'] = [
                    {k: v for k, v in img.items() if k != 'pil_image'}
                    for img in candidate_copy.get('images', [])
                ]
            serializable_data.append(candidate_copy)
        
        json_file = f"processed_resumes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(serializable_data, f, indent=2, ensure_ascii=False)
        
        print(f"Data saved: {json_file}")


def main():
    """Main execution function."""
    print("Resume Parser and ATS System")
    print("=" * 60)
    
    # Initialize parser
    parser = ResumeParser()
    
    # Process resumes
    resume_folder = input("Enter resume folder path: ").strip() or "./resumes"
    parser.process_resume_folder(resume_folder)
    
    if not parser.resume_data:
        print("No resumes processed")
        return
    
    # Job description
    job_description = """
    Looking for a Python Developer with machine learning and data science experience.
    Required: Python, pandas, scikit-learn, Django, REST APIs, SQL.
    Preferred: AWS, cloud platforms, strong problem-solving skills.
    """
    
    print("\nJob Description:")
    print(job_description)
    print("=" * 60)
    
    # Rank candidates
    top_candidates = parser.rank_candidates(job_description, top_k=10)
    
    # Generate report
    report_df = parser.generate_ats_report(top_candidates)
    
    # Display top candidates
    print("\nTOP CANDIDATES:")
    print("=" * 60)
    for i, candidate in enumerate(top_candidates[:5], 1):
        print(f"{i}. {candidate.get('name', 'N/A')} (Score: {candidate.get('total_score', 0):.3f})")
        print(f"   Email: {candidate.get('email', 'N/A')}")
        print(f"   Phone: {candidate.get('phone', 'N/A')}")
        print(f"   Skills: {', '.join(candidate.get('skills', [])[:5])}")
        print(f"   Profile Picture: {'Yes' if candidate.get('has_profile_picture') else 'No'}")
        print(f"   File: {candidate.get('filename', 'N/A')}")
        print("-" * 40)
    
    # Create dashboard
    parser.create_analytics_dashboard(report_df)
    
    # Save data
    parser.save_processed_data()
    
    print("\nProcessing complete. Check generated files for detailed analysis.")


if __name__ == "__main__":
    main()
